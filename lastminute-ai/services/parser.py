import pytesseract

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

"""
services/parser.py — Multi-format document parser
Handles: PDF, DOCX, PNG/JPG (handwritten notes via Tesseract OCR), TXT
"""

import os
import io
from pathlib import Path
from typing import List, Dict, Any
import fitz                        # PyMuPDF
import docx
from PIL import Image
import pytesseract

from config import get_settings

settings = get_settings()


class ParsedDocument:
    def __init__(self, filename: str, file_type: str):
        self.filename = filename
        self.file_type = file_type
        self.pages: List[Dict[str, Any]] = []   # [{page_num, text, source_label}]

    @property
    def full_text(self) -> str:
        return "\n\n".join(p["text"] for p in self.pages if p["text"].strip())

    @property
    def page_count(self) -> int:
        return len(self.pages)


def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """
    Route to the right parser based on file extension.
    Returns a ParsedDocument with per-page text and metadata.
    """
    ext = Path(filename).suffix.lower().lstrip(".")
    doc = ParsedDocument(filename=filename, file_type=ext)

    if ext == "pdf":
        _parse_pdf(file_bytes, doc)
    elif ext == "docx":
        _parse_docx(file_bytes, doc)
    elif ext in ("png", "jpg", "jpeg", "webp", "bmp"):
        _parse_image(file_bytes, doc)
    elif ext == "txt":
        _parse_txt(file_bytes, doc)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

    return doc


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _parse_pdf(file_bytes: bytes, doc: ParsedDocument):
    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        text = page.get_text("text").strip()

        # If a PDF page has very little text it's likely scanned — run OCR on it
        if len(text) < 50:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img).strip()

        if text:
            doc.pages.append({
                "page_num": page_num + 1,
                "text": text,
                "source_label": f"{doc.filename} — Page {page_num + 1}",
            })
    pdf.close()


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _parse_docx(file_bytes: bytes, doc: ParsedDocument):
    word_doc = docx.Document(io.BytesIO(file_bytes))
    full_text = "\n".join(
        para.text.strip()
        for para in word_doc.paragraphs
        if para.text.strip()
    )
    # Also extract tables
    for table in word_doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            full_text += "\n" + row_text

    # Treat DOCX as a single "page"
    if full_text.strip():
        doc.pages.append({
            "page_num": 1,
            "text": full_text.strip(),
            "source_label": f"{doc.filename}",
        })


# ─── IMAGE (handwritten notes via OCR) ───────────────────────────────────────

def _parse_image(file_bytes: bytes, doc: ParsedDocument):
    image = Image.open(io.BytesIO(file_bytes))

    # Preprocess for better OCR accuracy
    image = image.convert("L")          # Grayscale
    image = _upscale_if_small(image)

    custom_config = r"--oem 3 --psm 6"
    text = pytesseract.image_to_string(image, config=custom_config).strip()

    if text:
        doc.pages.append({
            "page_num": 1,
            "text": text,
            "source_label": f"{doc.filename} (handwritten notes)",
        })


def _upscale_if_small(img: Image.Image) -> Image.Image:
    """Upscale small images so Tesseract performs better."""
    w, h = img.size
    if w < 1200:
        scale = 1200 / w
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    return img


# ─── TXT ──────────────────────────────────────────────────────────────────────

def _parse_txt(file_bytes: bytes, doc: ParsedDocument):
    text = file_bytes.decode("utf-8", errors="ignore").strip()
    if text:
        doc.pages.append({
            "page_num": 1,
            "text": text,
            "source_label": f"{doc.filename}",
        })
